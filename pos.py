from tkinter import *
from tkinter.ttk import Label, Button, Entry
from tkinter import filedialog, messagebox
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import sqlite3
from PIL import Image, ImageTk
from io import BytesIO


con = sqlite3.connect("foodlist.db")
cur = con.cursor()
cur.execute("CREATE TABLE IF NOT EXISTS food(id INTEGER PRIMARY KEY, name TEXT, price REAL)")

root = Tk()
root.title('Python Point Of Sale')

itemDict = {}
amount = 0

def openAddItem():
     
    # Toplevel object which will 
    # be treated as a new window
    newWindow = Toplevel(root)

    titleLabel = Label(newWindow, text="Add a new food item!")
    titleLabel.grid(row=0, column=1, columnspan=2, pady=10)
 
    # sets the title of the
    # Toplevel widget
    newWindow.title("Add a new food item")
    nameLabel = Label(newWindow, text="Item name: ")
    nameLabel.grid(row=1, column=0, padx=5, pady=5, sticky="e")

    nameEntry = Entry(newWindow)
    nameEntry.grid(row=1, column=2, padx=5, pady=5, sticky="w")

    priceLabel = Label(newWindow, text="Price: ")
    priceLabel.grid(row=2, column=0, padx=5, pady=5, sticky="e")

    priceEntry = Entry(newWindow)
    priceEntry.grid(row=2, column=2, padx=5, pady=5, sticky="w")

    submitBtnFood = Button(newWindow, text="Submit", command=lambda: insertSQL(nameEntry.get().strip(), priceEntry.get().strip()))
    submitBtnFood.grid(row=3, column=1, columnspan=2, pady=10)

 
    # sets the geometry of toplevel
    newWindow.geometry("400x400")

def insertSQL(foodName, foodPrice):

    try:
        foodPrice = float(foodPrice)
    except ValueError:
        print("Invalid price format")
        return

    isAddedAlready = cur.execute("SELECT * FROM food WHERE name=?", (foodName,))
    if isAddedAlready:
        messagebox.showwarning("Unable to add item.", "Item is already added!")
    else:
        cur.execute("SELECT COUNT(id) FROM food")
        count = cur.fetchone()[0]
        cur.execute(f"INSERT INTO food(id, name, price) VALUES (?, ?, ?)", (count + 1, foodName, foodPrice))
        con.commit()

        # For debug purposes
        food = cur.execute("SELECT * FROM food ORDER BY id")
        for item in food:
            print(item)

def on_entry_click(event):
   if cashReceived.get() == "e.g. 200":
      cashReceived.delete(0, END)
      cashReceived.configure(foreground="black")

def on_focus_out(event):
   if cashReceived.get() == "":
      cashReceived.insert(0, "e.g. 200")
      cashReceived.configure(foreground="gray")

def update_total_label():
    totalAmount.config(text=f"Total amount: ${amount:.2f}")

def insertOrder(name:str, price:float):
    global amount
    if name in itemDict:
        itemDict[name] = round((itemDict[name] + price),2)
    else:
        itemDict[name] = price
    amount += price
    update_total_label()
    return itemDict

def generateReceipt(itemDict, save_path="Receipt.docx"):
    global amount
    amount = 0
    update_total_label()

    cash = cashReceived.get().strip()
    try:
        cash = float(cash)
        document = docx.Document()

        # Centered, bold "Store Name" header
        receiptTitle = document.add_paragraph('Shop Name')
        receiptTitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in receiptTitle.runs:
            run.bold = False
            run.font.size = Pt(48)

        a = document.add_paragraph('\n----------------------------------------------------------------------------------------------------------------------\n')
        a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        receiptTitle = document.add_paragraph('Receipt')
        receiptTitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in receiptTitle.runs:
            run.bold = False
            run.font.size = Pt(24)

        b = document.add_paragraph('\n----------------------------------------------------------------------------------------------------------------------\n')
        b.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Function to add an item with name on the left and price on the right
        def add_item(name, price):
            # Create a table with a single row and two columns
            table = document.add_table(rows=1, cols=2)
            
            # Set the name in the left cell
            cell_name = table.cell(0, 0)
            cell_name.text = name
            
            # Set the price in the right cell
            cell_price = table.cell(0, 1)
            cell_price.text = f"${price:.2f}"  # Format price to two decimal places
            
            # Align the text in each cell
            cell_name.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            cell_price.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        totalPrice = 0
        # Add items to the document 
        for item in itemDict:
            add_item(item, float(itemDict[item]))
            totalPrice += itemDict[item]

        table2 = document.add_table(rows=1, cols=2)
        cellText = table2.cell(0,0)
        cellText.text = "Total Amount"
        cellPrice = table2.cell(0, 1)
        cellPrice.text = f"${totalPrice:.2f}"

        cellText.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        cellPrice.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        c = document.add_paragraph('\n----------------------------------------------------------------------------------------------------------------------\n')
        c.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table3 = document.add_table(rows=2, cols=2)
        cellText2 = table3.cell(0,0)
        cellText2.text = "Cash Received"
        cellPrice2 = table3.cell(0, 1)
        cellPrice2.text = f"${cash}"
        cellText3 = table3.cell(1,0)
        cellText3.text = "Change"
        cellPrice3 = table3.cell(1, 1)
        cellPrice3.text = f"${(cash-totalPrice):.2f}"

        cellText2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        cellPrice2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        cellText3.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        cellPrice3.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        e = document.add_paragraph('\n----------------------------------------------------------------------------------------------------------------------\n')
        e.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        receiptTitle = document.add_paragraph('THANK YOU')
        receiptTitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in receiptTitle.runs:
            run.bold = True
            run.font.size = Pt(24)

        # Save the document
        document.save(save_path)
    except:
        messagebox.showwarning("Check Out fail", "Invalid cash format! (Numeric only!)")

mainLabel = Label(root, text="Point Of Sale (POS) System")
mainLabel.config(font=('Helvetica', 18, 'bold'))
mainLabel.pack()

desc = Label(root, text="Made with love by Mutton9558 💖")
desc.config(font=("Sans Serif", 10))
desc.pack(pady=(0,20))

separator = Frame(root, height=2, bd=1, relief="sunken")
separator.pack(fill="x", padx=5, pady=(0,15))

meeGoreng = Button(root, text="Noodles", command=lambda: insertOrder("Noodles", 5.90))
meeGoreng.pack(anchor="w", padx=(20))

totalAmount = Label(root, text=f"Total amount: ${amount:.2f}")
totalAmount.config(font=("Sans Serif", 10))
totalAmount.pack(pady=20)

cashReceived = Entry(root, width=25, foreground='gray')
cashReceived.insert(0, "e.g. 200")
cashReceived.bind("<FocusIn>", on_entry_click)
cashReceived.bind("<FocusOut>", on_focus_out)
cashReceived.pack()

checkOut = Button(root, text="Check Out", command=lambda: generateReceipt(itemDict))
checkOut.pack(anchor="s", pady=10)

btn = Button(root, text ="Click to add new food item!", command = openAddItem)
btn.pack(pady = 10)

root.mainloop()